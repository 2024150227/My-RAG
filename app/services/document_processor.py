# 文档处理器
import os
import hashlib
from pypdf import PdfReader
from openpyxl import load_workbook
from docx import Document
from app.config import settings
from app.utils.logger import logger

# PyMuPDF 是 PDF 抽图的工业标准（pypdf 抽图 API 残废）。
# 在没装的环境里允许 import 失败，extract_text_and_images 时再做兜底。
try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None
    logger.warning("PyMuPDF (fitz) 未安装，PDF 图片抽取将被跳过；pip install pymupdf 即可启用")


# 上传根目录（与 routes.py 的 uploads/<user_id>/ 保持一致）
UPLOAD_ROOT = "uploads"
# 每个文档的图片存放规则：uploads/<user_id>/_images/<doc_stem>/
IMAGES_SUBDIR = "_images"


def _safe_stem(filename: str) -> str:
    """把文件名（可能含中文/空格/标点）转成安全的目录名。

    取原文件名（不含扩展）的 md5 前 12 位，避免因为编码、长度、特殊字符
    在不同平台上踩坑。同一文件名命中同一 md5，天然幂等。
    """
    stem, _ = os.path.splitext(filename)
    safe = hashlib.md5(stem.encode('utf-8', errors='ignore')).hexdigest()[:12]
    return safe


class DocumentProcessor:
    def __init__(self):
        self.max_chunk_size = settings.MAX_CHUNK_SIZE
        self.overlap_ratio = settings.CHUNK_OVERLAP_RATIO

    def extract_text(self, file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        try:
            if ext == '.txt' or ext == '.md':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            elif ext == '.pdf':
                reader = PdfReader(file_path)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
            elif ext in ('.xlsx', '.xls'):
                workbook = load_workbook(file_path)
                text = ''
                for sheet in workbook.sheetnames:
                    worksheet = workbook[sheet]
                    for row in worksheet.iter_rows(values_only=True):
                        row_text = ' '.join(str(cell) for cell in row if cell)
                        text += row_text + '\n'
                return text
            elif ext == '.docx':
                # Word 2007+ 文档：抽段落 + 标题 + 表格
                # 注：旧版 .doc（二进制 OLE）不支持，需要先用 Word/LibreOffice 另存为 .docx
                doc = Document(file_path)
                parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    # 标题样式 (Heading 1/2/3...) 转成 Markdown 的 #/##/###，
                    # 这样后续 semantic_chunking 能按标题切块
                    style = (para.style.name or '').lower()
                    if style.startswith('heading'):
                        # 取末尾数字作为层级，没有就当 1 级
                        level = 1
                        for ch in style[::-1]:
                            if ch.isdigit():
                                level = max(1, min(6, int(ch)))
                                break
                        parts.append('#' * level + ' ' + text)
                    else:
                        parts.append(text)
                # 表格：每行用 ' | ' 拼接，非空行才保留
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        row_text = ' | '.join(c for c in cells if c)
                        if row_text:
                            parts.append(row_text)
                return '\n'.join(parts)
            else:
                logger.warning(f"不支持的文件格式: {ext}")
                return ''
        except Exception as e:
            logger.error(f"文件解析失败 {file_path}: {str(e)}")
            return ''

    def extract_images(self, file_path: str, user_id: str) -> list:
        """抽取文档中的内嵌图片，存盘并返回相对路径列表。

        返回值是 ``uploads/<user_id>/_images/<stem>/img_xxx.png`` 形式的
        相对路径，便于后面拼成 HTTP URL。

        - 仅支持 PDF / .docx；其他格式返回 []
        - 图片用 md5 命名，自然去重
        - 任何失败都吞掉（log warning），不影响文本入库
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        filename = os.path.basename(file_path)
        stem = _safe_stem(filename)

        out_dir = os.path.join(UPLOAD_ROOT, user_id, IMAGES_SUBDIR, stem)
        rel_paths: list[str] = []

        try:
            if ext == '.pdf':
                if fitz is None:
                    return []
                rel_paths = self._extract_pdf_images(file_path, out_dir)
            elif ext == '.docx':
                rel_paths = self._extract_docx_images(file_path, out_dir)
            # 其他格式不抽图
        except Exception as e:
            logger.warning(f"图片抽取失败（不影响文本入库） {file_path}: {str(e)}")
            return []

        if rel_paths:
            logger.info(f"文档 {filename} 抽出 {len(rel_paths)} 张图，存放于 {out_dir}")
        return rel_paths

    def _extract_pdf_images(self, file_path: str, out_dir: str) -> list:
        """PyMuPDF 遍历每页 get_images()，按 xref 抽图。"""
        os.makedirs(out_dir, exist_ok=True)
        rel_paths: list[str] = []
        seen_md5: set[str] = set()

        doc = fitz.open(file_path)
        try:
            for page_idx, page in enumerate(doc):
                images = page.get_images(full=True)
                for img_idx, img_info in enumerate(images):
                    xref = img_info[0]
                    try:
                        # extract_image 返回 dict: {'image': bytes, 'ext': 'png'/'jpeg'/...}
                        img = doc.extract_image(xref)
                    except Exception as e:
                        logger.warning(f"PDF 页 {page_idx} 第 {img_idx} 张图抽取失败: {e}")
                        continue
                    blob = img.get('image')
                    if not blob:
                        continue
                    # 太小的图（图标/分隔线）直接跳过
                    if len(blob) < 2048:
                        continue
                    md5 = hashlib.md5(blob).hexdigest()
                    if md5 in seen_md5:
                        continue
                    seen_md5.add(md5)

                    img_ext = (img.get('ext') or 'png').lower()
                    img_name = f"p{page_idx+1:03d}_{md5[:8]}.{img_ext}"
                    abs_path = os.path.join(out_dir, img_name)
                    if not os.path.exists(abs_path):
                        with open(abs_path, 'wb') as f:
                            f.write(blob)
                    rel_paths.append(abs_path.replace('\\', '/'))
        finally:
            doc.close()
        return rel_paths

    def _extract_docx_images(self, file_path: str, out_dir: str) -> list:
        """python-docx：遍历 part.related_parts，过滤 Image 类型。"""
        os.makedirs(out_dir, exist_ok=True)
        rel_paths: list[str] = []
        seen_md5: set[str] = set()

        doc = Document(file_path)
        # related_parts 是 dict: {rId: Part}
        for rel_id, part in doc.part.related_parts.items():
            content_type = getattr(part, 'content_type', '') or ''
            if not content_type.startswith('image/'):
                continue
            blob = part.blob
            if not blob or len(blob) < 2048:
                continue
            md5 = hashlib.md5(blob).hexdigest()
            if md5 in seen_md5:
                continue
            seen_md5.add(md5)

            # content_type 形如 'image/png' / 'image/jpeg' / 'image/x-emf'
            img_ext = content_type.split('/', 1)[-1].split(';', 1)[0].strip() or 'png'
            # emf/wmf 浏览器不显示，但先存着；前端可以过滤
            img_name = f"docx_{md5[:8]}.{img_ext}"
            abs_path = os.path.join(out_dir, img_name)
            if not os.path.exists(abs_path):
                with open(abs_path, 'wb') as f:
                    f.write(blob)
            rel_paths.append(abs_path.replace('\\', '/'))
        return rel_paths

    def extract_text_and_images(self, file_path: str, user_id: str) -> dict:
        """一次性抽出文本 + 图片（推荐入口）。

        返回 ``{"text": str, "images": [相对路径列表]}``。
        即使图片抽取失败，text 仍然保证返回；保持现有上传链路不被打断。
        """
        text = self.extract_text(file_path)
        images = self.extract_images(file_path, user_id)
        return {"text": text, "images": images}

    @staticmethod
    def _is_table_separator(line: str) -> bool:
        """判断是否为 Markdown 表格分隔行（|----|、|:---:| 等）。

        只含 | : - 空格 四种字符的行即为分隔行，直接跳过。
        """
        stripped = line.strip()
        return stripped.startswith('|') and all(ch in '|:- ' for ch in stripped)

    def semantic_chunking(self, content: str) -> list:
        chunks = []
        lines = content.split('\n')

        current_chunk = ""

        for line in lines:
            stripped = line.strip()

            # ── 表格分隔行直接跳过（不发散到 current_chunk） ──
            if self._is_table_separator(line):
                continue

            # ── 表格行：以 | 开头 → 整表保持在同一切块内 ──
            if stripped.startswith('|'):
                current_chunk += line + "\n"
                continue

            # ── 标题行：切一刀 ──
            if line.startswith('#'):
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = line + "\n"

            # ── 空行：切一刀（表格行已被上面的 continue 截住，不会误切表格） ──
            elif stripped == "":
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = ""

            # ── 普通文本：超过长度才切 ──
            else:
                if len(current_chunk) + len(line) > self.max_chunk_size:
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                        overlap_length = int(len(current_chunk) * self.overlap_ratio)
                        overlap_text = current_chunk[-overlap_length:] if overlap_length > 0 else ""
                        current_chunk = overlap_text + line + "\n"
                    else:
                        current_chunk = line + "\n"
                else:
                    current_chunk += line + "\n"

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return [chunk for chunk in chunks if len(chunk.strip()) > 20]

document_processor = DocumentProcessor()
